# -*- coding: utf-8 -*-
from __future__ import annotations

import calendar
import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
import threading
import time
import unicodedata
from logging.handlers import RotatingFileHandler
from copy import deepcopy
from datetime import date, datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Iterable
from uuid import uuid4

from models.contract_models import ContractTemplateProfile

logger = logging.getLogger(__name__)

_CONTRACT_LOGGER_LOCK = threading.Lock()
_CONTRACT_LOGGER_READY = False
_CONTRACT_LOG_FILE: Path | None = None

_PLACEHOLDER_RE = re.compile(
    r"x{3}\.x{3}\.x{3}-x{2}|xx\.xxx-xxx|xx/xx/xxxx|\(\d{2}\)\s*x{5}-x{4}|x{3}",
    re.IGNORECASE,
)
_NAMED_PLACEHOLDER_RE = re.compile(
    r"\{\{\s*([^{}]{1,120}?)\s*\}\}|\{\s*([^{}]{1,120}?)\s*\}|\[\s*([^\[\]]{1,160}?)\s*\]"
)
_CURLY_PLACEHOLDER_RE = re.compile(r"\{\{\s*[^{}]{1,120}\s*\}\}|\{\s*[^{}]{1,120}\s*\}")
_BRACKET_PLACEHOLDER_RE = re.compile(r"\[([^\[\]\n]{1,160})\]")

_WORD_CONVERT_LOCK = threading.RLock()
_TEMPLATE_CACHE_LOCK = threading.Lock()
_TEMPLATE_BYTES_CACHE: dict[str, tuple[float, bytes]] = {}

_UNIDADES = [
    "zero", "um", "dois", "tres", "quatro", "cinco", "seis", "sete", "oito", "nove",
    "dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove",
]
_DEZENAS = {
    20: "vinte", 30: "trinta", 40: "quarenta", 50: "cinquenta",
    60: "sessenta", 70: "setenta", 80: "oitenta", 90: "noventa",
}
_CENTENAS = {
    100: "cem", 200: "duzentos", 300: "trezentos", 400: "quatrocentos",
    500: "quinhentos", 600: "seiscentos", 700: "setecentos", 800: "oitocentos", 900: "novecentos",
}


def normalize_contract_type(raw: str) -> str:
    txt = (raw or "").strip().lower()
    plain = "".join(ch for ch in unicodedata.normalize("NFKD", txt) if not unicodedata.combining(ch))
    if "pix" in plain:
        return "pix"
    if "boleto" in plain:
        return "boleto"
    if "recepcao" in plain or "recep" in plain:
        return "recepcao"
    return plain


def normalize_contract_operation(raw: str) -> str:
    txt = (raw or "").strip().lower()
    plain = "".join(ch for ch in unicodedata.normalize("NFKD", txt) if not unicodedata.combining(ch))
    plain = re.sub(r"[^a-z0-9_]+", "_", plain)
    plain = re.sub(r"_+", "_", plain).strip("_")
    return plain or "padrao"


def resolve_contract_template(contract_type: str, operation: str = "padrao") -> Path:
    tipo = normalize_contract_type(contract_type)
    op = normalize_contract_operation(operation)
    profile = build_contract_template_profile(tipo, op)
    for candidate in profile.candidates:
        if candidate.exists():
            return candidate
    return profile.candidates[-1]


def build_contract_template_profile(contract_type: str, operation: str = "padrao") -> ContractTemplateProfile:
    tipo = normalize_contract_type(contract_type)
    op = normalize_contract_operation(operation)
    base = _contracts_dir()
    candidates = (
        base / f"contrato_{tipo}_{op}.docx",
        base / f"contrato_{op}_{tipo}.docx",
        base / f"contrato_{tipo}.docx",
    )
    return ContractTemplateProfile(contract_type=tipo, operation=op, candidates=candidates)


def _resolve_contract_logs_dir() -> Path:
    root_logger = logging.getLogger()
    for handler in list(getattr(root_logger, "handlers", []) or []):
        base = getattr(handler, "baseFilename", None)
        if not base:
            continue
        try:
            path = Path(str(base)).resolve().parent
            path.mkdir(parents=True, exist_ok=True)
            return path
        except Exception:
            continue
    fallback = Path.cwd() / "logs"
    fallback.mkdir(parents=True, exist_ok=True)
    return fallback


def _ensure_contract_file_logger():
    global _CONTRACT_LOGGER_READY, _CONTRACT_LOG_FILE
    if _CONTRACT_LOGGER_READY:
        return
    with _CONTRACT_LOGGER_LOCK:
        if _CONTRACT_LOGGER_READY:
            return
        logs_dir = _resolve_contract_logs_dir()
        log_file = logs_dir / "contract_generation.log"
        target = str(log_file.resolve())

        for h in list(getattr(logger, "handlers", []) or []):
            base = getattr(h, "baseFilename", None)
            if base and str(Path(str(base)).resolve()) == target:
                _CONTRACT_LOG_FILE = log_file
                _CONTRACT_LOGGER_READY = True
                return

        handler = RotatingFileHandler(
            log_file,
            maxBytes=2 * 1024 * 1024,
            backupCount=5,
            encoding="utf-8",
        )
        handler.setLevel(logging.INFO)
        handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] [contract] %(message)s"))
        logger.addHandler(handler)

        _CONTRACT_LOG_FILE = log_file
        _CONTRACT_LOGGER_READY = True
        logger.info("Contrato: log dedicado ativo (%s).", str(log_file))


def generate_contract_pdf(
    cliente: dict,
    dependentes: list[dict],
    contract_type: str,
    operation: str = "padrao",
    output_dir: Path | None = None,
) -> Path:
    _ensure_contract_file_logger()
    started_at = time.perf_counter()
    tipo = normalize_contract_type(contract_type)
    if tipo not in {"pix", "boleto", "recepcao"}:
        raise ValueError("Tipo de contrato inválido. Use pix, boleto ou recepcao.")
    _validate_contract_inputs(cliente, dependentes, tipo, operation)

    template = resolve_contract_template(tipo, operation)
    if not template.exists():
        raise FileNotFoundError(
            f"Template não encontrado para tipo={tipo} e operacao={normalize_contract_operation(operation)}: {template}"
        )
    logger.info("Contrato: template selecionado (%s).", str(template))

    try:
        Document, Paragraph = _import_docx()
        doc = _load_template_doc(Document, template)

        deps_normalizados = _normalize_dependentes_for_contract(dependentes)
        payload = _build_payload(cliente, deps_normalizados)
        _fill_main_fields(doc, payload)
        applied_dep_block = _apply_dependentes_block(doc, payload["dependentes"], Paragraph)
        if payload["dependentes"] and not applied_dep_block:
            logger.warning(
                "Contrato: bloco de dependentes não encontrado no template (%s). Dependentes podem não aparecer.",
                str(template),
            )
        _assert_no_placeholders(doc)

        out_dir = Path(output_dir) if output_dir else _default_downloads_dir()
        out_dir.mkdir(parents=True, exist_ok=True)

        hoje_iso = datetime.now().strftime("%Y-%m-%d")
        nome_slug = _slugify_filename(payload["nome"] or "cliente")
        target_pdf = _unique_path(out_dir / f"contrato_{tipo}_{nome_slug}_{hoje_iso}.pdf")

        tmp_docx = Path(tempfile.gettempdir()) / f"medcontract_contract_{uuid4().hex}.docx"
        tmp_pdf = tmp_docx.with_suffix(".pdf")
        try:
            doc.save(str(tmp_docx))
            _convert_docx_to_pdf(tmp_docx, target_pdf)
        finally:
            for p in (tmp_docx, tmp_pdf):
                try:
                    if p.exists():
                        p.unlink()
                except Exception:
                    pass

        elapsed_ms = int((time.perf_counter() - started_at) * 1000)
        logger.info("Contrato: PDF gerado (%s) em %sms.", str(target_pdf), elapsed_ms)
        return target_pdf
    except Exception:
        elapsed_ms = int((time.perf_counter() - started_at) * 1000)
        logger.exception("Contrato: falha na geração (tipo=%s, operacao=%s, tempo=%sms).", tipo, normalize_contract_operation(operation), elapsed_ms)
        raise


def _import_docx():
    try:
        from docx import Document  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "Biblioteca python-docx não instalada. Instale com: pip install python-docx"
        ) from exc
    return Document, Paragraph


def _contracts_dir() -> Path:
    base = Path(__file__).resolve().parents[1]
    for folder in ("Contratos", "contratos"):
        p = base / folder
        if p.exists():
            return p
    return base / "Contratos"


def _load_template_doc(Document, template_path: Path):
    path = template_path.resolve()
    key = str(path)
    mtime = path.stat().st_mtime

    with _TEMPLATE_CACHE_LOCK:
        cached = _TEMPLATE_BYTES_CACHE.get(key)
        if cached and cached[0] == mtime:
            template_bytes = cached[1]
        else:
            template_bytes = path.read_bytes()
            _TEMPLATE_BYTES_CACHE[key] = (mtime, template_bytes)

    return Document(io.BytesIO(template_bytes))


def _default_downloads_dir() -> Path:
    d = Path.home() / "Downloads"
    if d.exists():
        return d
    return Path.cwd()


def _unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem, suffix = path.stem, path.suffix
    idx = 2
    while True:
        candidate = path.with_name(f"{stem}_{idx}{suffix}")
        if not candidate.exists():
            return candidate
        idx += 1


def _slugify_filename(value: str) -> str:
    txt = "".join(ch for ch in unicodedata.normalize("NFKD", value or "") if not unicodedata.combining(ch))
    txt = re.sub(r"[^a-zA-Z0-9]+", "_", txt.strip().lower())
    txt = re.sub(r"_+", "_", txt).strip("_")
    return txt or "cliente"


def _plain_upper(value: str) -> str:
    txt = "".join(ch for ch in unicodedata.normalize("NFKD", value or "") if not unicodedata.combining(ch))
    return txt.upper()


def _normalize_placeholder_key(value: str) -> str:
    txt = "".join(ch for ch in unicodedata.normalize("NFKD", value or "") if not unicodedata.combining(ch))
    txt = txt.strip().lower()
    txt = re.sub(r"[^a-z0-9]+", "_", txt)
    txt = re.sub(r"_+", "_", txt).strip("_")
    return txt


def _format_forma_pagamento(value: str) -> str:
    tipo = normalize_contract_type(value)
    if tipo == "pix":
        return "PIX"
    if tipo == "recepcao":
        return "RECEPÇÃO"
    if tipo == "boleto":
        return "BOLETO"
    return (value or "").strip() or "-"


def _digits(value: str) -> str:
    return "".join(ch for ch in (value or "") if ch.isdigit())


def _mask_cpf_for_log(value: str) -> str:
    d = _digits(value)
    if len(d) >= 11:
        return f"{d[:3]}.***.***-{d[-2:]}"
    if d:
        return f"{d[:2]}***"
    return "-"


def _mask_email_for_log(value: str) -> str:
    txt = str(value or "").strip()
    if "@" not in txt:
        return "-"
    user, _, domain = txt.partition("@")
    if not user:
        return f"***@{domain}"
    return f"{user[0]}***@{domain}"


def _sanitize_text_for_doc(value: str) -> str:
    txt = _fix_common_mojibake(str(value or ""))
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt


def _normalize_dependentes_for_contract(dependentes: list[dict] | None) -> list[dict]:
    out: list[dict] = []
    seen_cpfs: set[str] = set()
    dropped_empty = 0
    dropped_dup = 0
    for dep in dependentes or []:
        item = dict(dep or {})
        nome = _sanitize_text_for_doc(item.get("nome") or "")
        cpf_digits = _digits(item.get("cpf") or "")
        cpf_fmt = _format_cpf(cpf_digits)
        data_nascimento = _iso_to_br_date(item.get("data_nascimento") or "")
        if not nome and (not cpf_digits) and data_nascimento in {"", "-"}:
            dropped_empty += 1
            continue
        if len(cpf_digits) == 11 and cpf_digits in seen_cpfs:
            dropped_dup += 1
            continue
        if len(cpf_digits) == 11:
            seen_cpfs.add(cpf_digits)
        out.append(
            {
                "nome": nome or "-",
                "cpf": cpf_fmt,
                "data_nascimento": data_nascimento or "-",
            }
        )
    if dropped_empty or dropped_dup:
        logger.warning(
            "Contrato: dependentes normalizados (removidos_vazios=%s, removidos_duplicados=%s, finais=%s).",
            dropped_empty,
            dropped_dup,
            len(out),
        )
    return out


def _validate_contract_inputs(cliente: dict, dependentes: list[dict], tipo: str, operation: str):
    if not isinstance(cliente, dict):
        raise ValueError("Dados do cliente inválidos para geração de contrato.")
    nome = _sanitize_text_for_doc(cliente.get("nome") or "")
    if not nome:
        raise ValueError("Nome do cliente é obrigatório para gerar o contrato.")
    cpf_digits = _digits(cliente.get("cpf") or "")
    if len(cpf_digits) != 11:
        raise ValueError("CPF do cliente inválido para geração do contrato.")
    if not isinstance(dependentes, list):
        raise ValueError("Lista de dependentes inválida para geração do contrato.")
    logger.info(
        "Contrato: entrada validada (tipo=%s, operacao=%s, cliente=%s, cpf=%s, email=%s, dependentes=%s).",
        tipo,
        normalize_contract_operation(operation),
        nome[:80],
        _mask_cpf_for_log(cpf_digits),
        _mask_email_for_log(cliente.get("email") or ""),
        len(dependentes),
    )


def _format_cpf(value: str) -> str:
    d = _digits(value)
    if len(d) == 11:
        return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}"
    return (value or "").strip() or "-"


def _format_cep(value: str) -> str:
    d = _digits(value)
    if len(d) == 8:
        return f"{d[:2]}.{d[2:5]}-{d[5:]}"
    return (value or "").strip() or "-"


def _format_phone(value: str) -> str:
    d = _digits(value)
    if len(d) == 11:
        return f"({d[:2]}) {d[2:7]}-{d[7:]}"
    if len(d) == 10:
        return f"({d[:2]}) {d[2:6]}-{d[6:]}"
    return (value or "").strip() or "-"


def _iso_to_br_date(value: str) -> str:
    s = (value or "").strip()
    if not s:
        return "-"
    if "-" in s and len(s) >= 10:
        try:
            y, m, d = s[:10].split("-")
            return f"{d}/{m}/{y}"
        except Exception:
            return s
    if "/" in s:
        return s
    return s


def _fix_common_mojibake(value: str) -> str:
    txt = str(value or "")
    if not txt:
        return ""
    fixes = {
        "â€¢": "•",
        "â€“": "-",
        "â€”": "-",
        "NÂº": "Nº",
        "Âº": "º",
    }
    for bad, good in fixes.items():
        txt = txt.replace(bad, good)
    return txt.strip()


def _build_address(endereco_raw: str, cep_raw: str) -> str:
    endereco_raw = _sanitize_text_for_doc(endereco_raw)
    # Alguns ambientes convertem bullets para "?".
    endereco_raw = re.sub(r"\s+\?{1,3}\s+", " • ", endereco_raw)
    endereco_raw = endereco_raw.replace("|", " • ")

    numero = "S/N"
    numero_match = re.search(
        r"\bN(?:Â)?(?:[º°o]|[úu]mero|\?+)?\.?\s*[:.]?\s*([A-Za-z0-9\-\/]+)\b",
        endereco_raw,
        flags=re.IGNORECASE,
    )
    if numero_match:
        numero = numero_match.group(1).strip() or "S/N"
        endereco_raw = (endereco_raw[:numero_match.start()] + " " + endereco_raw[numero_match.end():]).strip()

    parts = [p.strip(" ,.-") for p in re.split(r"\s*(?:•|·|∙|â€¢|,)\s*", endereco_raw) if p.strip(" ,.-")]
    if len(parts) < 2:
        parts = [p.strip(" ,.-") for p in re.split(r"\s+-\s+", endereco_raw) if p.strip(" ,.-")]

    sem_numero: list[str] = []
    for p in parts:
        p_norm = p.upper().strip()
        if p_norm in {"RJ", "RJ.", "R.J."}:
            continue
        if re.match(r"^N(?:Â)?(?:[º°o]|[úu]mero|\?+)?\.?\s*$", p, flags=re.IGNORECASE):
            continue
        sem_numero.append(p)

    logradouro = sem_numero[0] if len(sem_numero) > 0 else "-"
    bairro = sem_numero[1] if len(sem_numero) > 1 else "-"
    cidade = sem_numero[2] if len(sem_numero) > 2 else "RIO DE JANEIRO"
    cep = _format_cep(cep_raw)

    return f"{logradouro}, Nº {numero}, {bairro} - {cidade} - CEP.: {cep}, RIO DE JANEIRO."


def _money_number_br(value: float) -> str:
    q = Decimal(str(value or 0)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    s = f"{q:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return s


def _extenso_ate_999(n: int) -> str:
    if n < 20:
        return _UNIDADES[n]
    if n < 100:
        d = (n // 10) * 10
        r = n % 10
        return _DEZENAS[d] if r == 0 else f"{_DEZENAS[d]} e {_UNIDADES[r]}"
    if n == 100:
        return "cem"
    c = (n // 100) * 100
    r = n % 100
    c_txt = _CENTENAS.get(c, "")
    return c_txt if r == 0 else f"{c_txt} e {_extenso_ate_999(r)}"


def _numero_extenso(n: int) -> str:
    if n == 0:
        return "zero"
    if n < 0:
        return f"menos {_numero_extenso(abs(n))}"

    grupos: list[int] = []
    while n > 0:
        grupos.append(n % 1000)
        n //= 1000

    escalas = [
        ("", ""),
        ("mil", "mil"),
        ("milhao", "milhoes"),
        ("bilhao", "bilhoes"),
    ]

    partes: list[str] = []
    for idx in range(len(grupos) - 1, -1, -1):
        g = grupos[idx]
        if g == 0:
            continue
        if idx == 0:
            partes.append(_extenso_ate_999(g))
            continue
        sing, plur = escalas[idx]
        if idx == 1 and g == 1:
            partes.append("mil")
            continue
        ext = _extenso_ate_999(g)
        escala = sing if g == 1 else plur
        partes.append(f"{ext} {escala}".strip())

    if not partes:
        return "zero"
    if len(partes) == 1:
        return partes[0]
    return ", ".join(partes[:-1]) + " e " + partes[-1]


def _money_extenso(value: float) -> str:
    q = Decimal(str(value or 0)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    inteiro = int(q)
    centavos = int((q - Decimal(inteiro)) * 100)

    partes: list[str] = []
    if inteiro > 0:
        real_label = "real" if inteiro == 1 else "reais"
        partes.append(f"{_numero_extenso(inteiro)} {real_label}")
    if centavos > 0:
        cent_label = "centavo" if centavos == 1 else "centavos"
        partes.append(f"{_numero_extenso(centavos)} {cent_label}")
    if not partes:
        return "zero real"
    return " e ".join(partes)


def _safe_int(value, default: int = 0) -> int:
    try:
        return int(value)
    except Exception:
        return int(default)


def _safe_float(value, default: float = 0.0) -> float:
    try:
        if value is None:
            return float(default)
        txt = str(value).strip()
        if not txt:
            return float(default)
        # aceita formato BR/US, removendo separador de milhar
        txt = txt.replace(".", "").replace(",", ".") if "," in txt else txt
        return float(txt)
    except Exception:
        return float(default)


def _first_payment_date(today: date, vencimento_dia: int) -> date:
    day = max(1, min(31, _safe_int(vencimento_dia, 10)))
    y, m = today.year, today.month
    month_last = calendar.monthrange(y, m)[1]
    candidate = date(y, m, min(day, month_last))
    if candidate < today:
        if m == 12:
            y, m = y + 1, 1
        else:
            m += 1
        month_last = calendar.monthrange(y, m)[1]
        candidate = date(y, m, min(day, month_last))
    return candidate


def _build_payload(cliente: dict, dependentes: list[dict]) -> dict:
    hoje = datetime.now().date()
    nome = _sanitize_text_for_doc(cliente.get("nome") or "")
    cpf = _format_cpf(cliente.get("cpf", ""))
    data_nascimento = _iso_to_br_date(cliente.get("data_nascimento", ""))
    telefone = _format_phone(cliente.get("telefone", ""))
    email = _sanitize_text_for_doc(cliente.get("email") or "") or "-"
    endereco = _build_address(cliente.get("endereco", ""), cliente.get("cep", ""))

    valor = _safe_float(cliente.get("valor_mensal"), 0.0)
    valor_num = _money_number_br(valor)
    valor_extenso = _money_extenso(valor)

    venc_dia = max(1, min(31, _safe_int(cliente.get("vencimento_dia", 10), 10)))
    venc_extenso = _numero_extenso(venc_dia)
    primeiro_pag = _first_payment_date(hoje, venc_dia).strftime("%d/%m/%Y")
    forma_pagamento = _format_forma_pagamento(cliente.get("forma_pagamento", ""))

    deps_out: list[dict] = []
    for d in dependentes or []:
        deps_out.append({
            "nome": _sanitize_text_for_doc(d.get("nome") or "") or "-",
            "cpf": _format_cpf(d.get("cpf", "")),
            "data_nascimento": _iso_to_br_date(d.get("data_nascimento", "")),
        })

    return {
        "nome": nome,
        "cpf": cpf,
        "data_nascimento": data_nascimento,
        "endereco": endereco,
        "telefone": telefone,
        "email": email,
        "data_adesao": hoje.strftime("%d/%m/%Y"),
        "valor_mensal": f"R$ {valor_num}",
        "valor_num": valor_num,
        "valor_extenso": valor_extenso,
        "dia_vencimento": str(venc_dia),
        "venc_dia": str(venc_dia),
        "venc_extenso": venc_extenso,
        "dia_por_extenso": venc_extenso,
        "primeiro_pagamento": primeiro_pag,
        "forma_pagamento": forma_pagamento,
        "data_assinatura": hoje.strftime("%d/%m/%Y"),
        # Placeholder contratual de limite adicional ao IPCA.
        "reajuste_maximo_pct": str(cliente.get("reajuste_maximo_pct") or "10"),
        "dependentes": deps_out,
    }


def _iter_all_paragraphs(doc) -> Iterable:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for t2 in cell.tables:
                    for r2 in t2.rows:
                        for c2 in r2.cells:
                            for p2 in c2.paragraphs:
                                yield p2


def _paragraph_text(paragraph) -> str:
    if getattr(paragraph, "runs", None):
        return "".join(r.text for r in paragraph.runs)
    return paragraph.text or ""


def _replace_span_in_runs(paragraph, start: int, end: int, replacement: str) -> bool:
    if start < 0 or end < start:
        return False
    runs = list(getattr(paragraph, "runs", []))
    if not runs:
        return False

    pos = 0
    start_ref = None
    end_ref = None

    for idx, run in enumerate(runs):
        txt = run.text or ""
        next_pos = pos + len(txt)
        if start_ref is None and start < next_pos:
            start_ref = (idx, start - pos)
        if end_ref is None and end <= next_pos:
            end_ref = (idx, end - pos)
            break
        pos = next_pos

    if start_ref is None:
        return False
    if end_ref is None:
        end_ref = (len(runs) - 1, len(runs[-1].text or ""))

    si, so = start_ref
    ei, eo = end_ref

    if si == ei:
        old = runs[si].text or ""
        runs[si].text = old[:so] + replacement + old[eo:]
        return True

    start_txt = runs[si].text or ""
    end_txt = runs[ei].text or ""
    runs[si].text = start_txt[:so] + replacement + end_txt[eo:]
    for j in range(si + 1, ei + 1):
        runs[j].text = ""
    return True


def _replace_first_regex(paragraph, pattern: str, replacement: str, flags=re.IGNORECASE) -> bool:
    txt = _paragraph_text(paragraph)
    m = re.search(pattern, txt, flags=flags)
    if not m:
        return False
    return _replace_span_in_runs(paragraph, m.start(), m.end(), replacement)


def _replace_after_label(paragraph, label: str, placeholder_pattern: str, replacement: str) -> bool:
    txt = _paragraph_text(paragraph)
    txt_plain = _plain_upper(txt)
    label_plain = _plain_upper(label)
    idx = txt_plain.find(label_plain)
    if idx < 0:
        return False
    start_search = idx + len(label)
    m = re.search(placeholder_pattern, txt[start_search:], flags=re.IGNORECASE)
    if not m:
        return False
    return _replace_span_in_runs(
        paragraph,
        start_search + m.start(),
        start_search + m.end(),
        replacement,
    )


def _replace_tail_after_colon(paragraph, replacement: str) -> bool:
    txt = _paragraph_text(paragraph)
    if ":" not in txt:
        return False
    start = txt.find(":") + 1
    return _replace_span_in_runs(paragraph, start, len(txt), f" {replacement}")


def _set_paragraph_text(paragraph, value: str):
    runs = list(getattr(paragraph, "runs", []) or [])
    if not runs:
        paragraph.text = value
        return
    runs[0].text = value
    for r in runs[1:]:
        r.text = ""


def _placeholder_values(payload: dict, *, include_dependentes: bool = True) -> dict[str, str]:
    values: dict[str, str] = {}

    def _put(keys: Iterable[str], value: str):
        txt = str(value or "")
        for key in keys:
            norm = _normalize_placeholder_key(key)
            if not norm:
                continue
            values[norm] = txt
            values[norm.replace("_", "")] = txt

    _put(("nome", "nome_completo_do_a_aderente", "nome_completo_do_aderente", "nome_titular"), payload.get("nome", ""))
    _put(("cpf",), payload.get("cpf", ""))
    _put(("data_nascimento", "datanascimento"), payload.get("data_nascimento", ""))
    _put(("endereco", "endereco_completo"), payload.get("endereco", ""))
    _put(("telefone", "celular"), payload.get("telefone", ""))
    _put(("email", "e-mail"), payload.get("email", ""))
    _put(("data_adesao", "dataadesao", "data_da_adesao"), payload.get("data_adesao", ""))
    _put(("valor_mensal", "valormensal"), payload.get("valor_mensal", ""))
    _put(("valor",), payload.get("valor_num", ""))
    _put(("valor_num", "valornum"), payload.get("valor_num", ""))
    _put(("valor_extenso", "valorextenso", "valor_por_extenso"), payload.get("valor_extenso", ""))
    _put(("dia_vencimento", "vencimento_dia", "dia"), payload.get("dia_vencimento", payload.get("venc_dia", "")))
    _put(("venc_dia", "vencimento"), payload.get("venc_dia", ""))
    _put(("venc_extenso", "dia_por_extenso"), payload.get("venc_extenso", ""))
    _put(("primeiro_pagamento", "data_do_primeiro_pagamento"), payload.get("primeiro_pagamento", ""))
    _put(
        (
            "forma_pagamento",
            "boleto_cartao_pix",
            "boleto_cartao_pix_chave_pix_cnpj_42_591_560_0001_08",
        ),
        payload.get("forma_pagamento", ""),
    )
    _put(("data_assinatura", "data_de_assinatura"), payload.get("data_assinatura", payload.get("data_adesao", "")))
    reajuste_pct = str(payload.get("reajuste_maximo_pct", "10") or "10").strip()
    if reajuste_pct and "%" not in reajuste_pct:
        reajuste_pct = f"{reajuste_pct}%"
    _put(("x", "x_pct"), reajuste_pct)

    if include_dependentes:
        deps = list(payload.get("dependentes", []) or [])
        for idx in range(1, 11):
            dep = dict(deps[idx - 1] or {}) if idx <= len(deps) else {}
            nome = str(dep.get("nome") or "").strip() or "-"
            cpf = str(dep.get("cpf") or "").strip() or "-"
            data_nasc = str(dep.get("data_nascimento") or "").strip() or "-"
            parentesco = str(dep.get("parentesco") or dep.get("grau_parentesco") or "").strip() or "Dependente"
            _put((f"dep{idx}_nome", f"dependente{idx}_nome", f"nome_do_dependente_{idx}"), nome)
            _put((f"dep{idx}_cpf", f"dependente{idx}_cpf", f"cpf_do_dependente_{idx}"), cpf)
            _put((f"dep{idx}_data_nascimento", f"dependente{idx}_data_nascimento"), data_nasc)
            _put((f"data_nascimento_dependente_{idx}", f"data_nasc_dependente_{idx}"), data_nasc)
            _put((f"dep{idx}_parentesco", f"dependente{idx}_parentesco"), parentesco)

    return values


def _replace_named_placeholders(paragraph, values: dict[str, str]):
    txt = _paragraph_text(paragraph)
    has_curly = ("{" in txt and "}" in txt)
    has_bracket = ("[" in txt and "]" in txt)
    if not has_curly and not has_bracket:
        return

    def _sub(match: re.Match) -> str:
        key_raw = (match.group(1) or match.group(2) or match.group(3) or "").strip()
        norm = _normalize_placeholder_key(key_raw)
        if not norm:
            return match.group(0)
        repl = values.get(norm)
        if repl is None:
            repl = values.get(norm.replace("_", ""))
        if repl is None and norm.startswith("boleto_") and "pix" in norm and "cnpj" in norm:
            repl = values.get("forma_pagamento")
        return repl if repl is not None else match.group(0)

    replaced = _NAMED_PLACEHOLDER_RE.sub(_sub, txt)
    if replaced != txt:
        _set_paragraph_text(paragraph, replaced)


def _replace_literal(paragraph, pattern: str, replacement: str, flags=re.IGNORECASE) -> bool:
    txt = _paragraph_text(paragraph)
    updated = re.sub(pattern, replacement, txt, flags=flags)
    if updated == txt:
        return False
    _set_paragraph_text(paragraph, updated)
    return True


def _fill_main_fields(doc, payload: dict):
    named_values = _placeholder_values(payload, include_dependentes=False)
    forma_pagamento = str(payload.get("forma_pagamento") or "").strip().upper()
    is_pix = "PIX" in forma_pagamento
    for p in _iter_all_paragraphs(doc):
        _replace_named_placeholders(p, named_values)
        txt = _paragraph_text(p)
        up = _plain_upper(txt)

        if "NOME COMPLETO:" in up:
            _replace_after_label(p, "NOME COMPLETO:", r"(x{3}|\[[^\]]+\])", payload["nome"])
        if "CPF:" in up and "DEPENDENTE" not in up:
            _replace_after_label(p, "CPF:", r"(x{3}\.x{3}\.x{3}-x{2}|\[[^\]]+\])", payload["cpf"])
        if "DATA DE NASCIMENTO:" in up and "NOME:" not in up:
            _replace_after_label(p, "DATA DE NASCIMENTO:", r"(xx/xx/xxxx|x{3}|\[[^\]]+\])", payload["data_nascimento"])
        if "ENDERECO RESIDENCIAL:" in up:
            _replace_tail_after_colon(p, payload["endereco"])
        if "TELEFONE:" in up:
            _replace_after_label(p, "TELEFONE:", r"(\(\d{2}\)\s*x{5}-x{4}|\[[^\]]+\])", payload["telefone"])
        if "E-MAIL:" in up:
            _replace_after_label(p, "E-MAIL:", r"(x{3}|\[[^\]]+\])", payload["email"])
        if "DATA DA ADESAO:" in up:
            _replace_after_label(p, "DATA DA ADESAO:", r"(x{3}|\[[^\]]+\])", payload["data_adesao"])
        if "VALOR MENSAL AJUSTADO NESSE CONTRATO:" in up:
            if not _replace_first_regex(
                p,
                r"R\$\s*xxx\s*\(xxx\)",
                f"R$ {payload['valor_num']}({payload['valor_extenso']})",
            ):
                _replace_first_regex(
                    p,
                    r"R\$\s*\[[^\]]+\]\s*\(\s*\[[^\]]+\]\s*\)",
                    f"R$ {payload['valor_num']} ({payload['valor_extenso']})",
                )
        if "VENCIMENTO TODO DIA" in up:
            if not _replace_after_label(
                p,
                "VENCIMENTO TODO DIA",
                r"xxx\s*\(xxx\)",
                f"{payload['venc_dia']} ({payload['venc_extenso']})",
            ):
                _replace_after_label(
                    p,
                    "VENCIMENTO TODO DIA",
                    r"(\[[^\]]+\]\s*\(\s*\[[^\]]+\]\s*\)|\[[^\]]+\])",
                    f"{payload['venc_dia']} ({payload['venc_extenso']})",
                )
        if "PRIMEIRO PAGAMENTO EM" in up:
            _replace_after_label(
                p,
                "PRIMEIRO PAGAMENTO EM",
                r"(x{3}|\[[^\]]+\])",
                payload["primeiro_pagamento"],
            )
            _replace_after_label(
                p,
                "PAGAMENTO",
                r"\[[^\]]+\]",
                payload["forma_pagamento"],
            )
        if is_pix:
            _replace_literal(
                p,
                r"\bPAGAMENTO\s+PIX\b",
                "PAGAMENTO VIA PIX (CHAVE CNPJ 42.591.560/0001-08)",
            )


def _apply_dependentes_block(doc, dependentes: list[dict], Paragraph) -> bool:
    paragraphs = doc.paragraphs
    if not paragraphs:
        return False

    start_idx = None
    end_idx = None
    header_txt = _plain_upper("O (A) ADERENTE, NESTE ATO INSCREVE OS SEGUINTES DEPENDENTES")
    stop_txt = _plain_upper("CONDICOES DOS SERVICOS OFERECIDOS AO(S) ADERENTE(S)")

    for i, p in enumerate(paragraphs):
        t = _plain_upper(_paragraph_text(p))
        if start_idx is None and header_txt in t:
            start_idx = i
            continue
        if start_idx is not None and stop_txt in t:
            end_idx = i
            break

    if start_idx is None:
        return False
    if end_idx is None:
        end_idx = min(len(paragraphs), start_idx + 4)
    if end_idx <= start_idx:
        return False

    block = doc.paragraphs[start_idx:end_idx]
    if not block:
        return False

    header_tpl = deepcopy(block[0]._p)
    dep_title_tpl = deepcopy(block[1]._p if len(block) > 1 else block[0]._p)
    dep_nome_tpl = deepcopy(block[2]._p if len(block) > 2 else block[-1]._p)
    dep_cpf_tpl = deepcopy(block[3]._p if len(block) > 3 else block[-1]._p)

    anchor = doc.paragraphs[end_idx] if end_idx < len(doc.paragraphs) else None

    for para in reversed(block):
        try:
            para._p.getparent().remove(para._p)
        except Exception:
            pass

    if not dependentes:
        return True

    def _insert_before_anchor(p_xml):
        if anchor is None:
            doc._body._element.append(p_xml)
            return Paragraph(p_xml, doc._body)
        anchor._p.addprevious(p_xml)
        return Paragraph(p_xml, anchor._parent)

    _insert_before_anchor(deepcopy(header_tpl))

    for i, dep in enumerate(dependentes, start=1):
        p_title = _insert_before_anchor(deepcopy(dep_title_tpl))
        _replace_first_regex(p_title, r"DEPENDENTE\s+\d+", f"DEPENDENTE {i}")
        _replace_first_regex(p_title, r"\[\s*NOME\s+DO\s+DEPENDENTE\s+\d+\s*\]", dep["nome"])

        p_nome = _insert_before_anchor(deepcopy(dep_nome_tpl))
        if not _replace_after_label(p_nome, "NOME:", r"(x{3}|\[[^\]]+\])", dep["nome"]):
            _replace_first_regex(p_nome, r"\[\s*NOME\s+DO\s+DEPENDENTE\s+\d+\s*\]", dep["nome"])

        p_cpf = _insert_before_anchor(deepcopy(dep_cpf_tpl))
        if not _replace_after_label(p_cpf, "CPF:", r"(x{3}\.x{3}\.x{3}-x{2}|x{3}|\[[^\]]+\])", dep["cpf"]):
            _replace_first_regex(p_cpf, r"\[\s*CPF\s+DO\s+DEPENDENTE\s+\d+\s*\]", dep["cpf"])
        if not _replace_after_label(
            p_cpf,
            "DATA DE NASCIMENTO:",
            r"(\d{2}/\d{2}/\d{4}|xx/xx/xxxx|x{3}|\[[^\]]+\])",
            dep["data_nascimento"],
        ):
            _replace_first_regex(
                p_cpf,
                r"\[\s*DATA\s+(?:NASCIMENTO|NASC)\s+DEPENDENTE\s+\d+\s*\]",
                dep["data_nascimento"],
            )
    return True


def _assert_no_placeholders(doc):
    hint_keys = {
        "nome",
        "cpf",
        "data",
        "endereco",
        "telefone",
        "email",
        "valor",
        "dia",
        "pagamento",
        "dependente",
        "assinatura",
        "boleto",
        "pix",
        "cartao",
        "x",
    }
    for p in _iter_all_paragraphs(doc):
        txt = _paragraph_text(p)
        if _PLACEHOLDER_RE.search(txt):
            preview = (txt or "").strip()[:220]
            raise RuntimeError(
                "Ainda existem placeholders 'xxx' no contrato. Revise o template e os dados. "
                f"Trecho: {preview}"
            )
        if _CURLY_PLACEHOLDER_RE.search(txt):
            preview = (txt or "").strip()[:220]
            raise RuntimeError(
                "Ainda existem placeholders com chaves no contrato. Revise o template e os dados. "
                f"Trecho: {preview}"
            )
        for m in _BRACKET_PLACEHOLDER_RE.finditer(txt):
            norm = _normalize_placeholder_key(m.group(1) or "")
            if not norm:
                continue
            if any(k in norm for k in hint_keys):
                preview = (txt or "").strip()[:220]
                raise RuntimeError(
                    "Ainda existem placeholders com colchetes no contrato. Revise o template e os dados. "
                    f"Trecho: {preview}"
                )


def _convert_via_word_com(src_docx: Path, dst_pdf: Path) -> tuple[bool, str | None]:
    word = None
    document = None
    com_ready = False
    keep_active = str(os.getenv("MEDCONTRACT_WORD_KEEP_ACTIVE", "1") or "1").strip().lower() in {
        "1", "true", "yes", "on"
    }
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore

        pythoncom.CoInitialize()
        com_ready = True

        with _WORD_CONVERT_LOCK:
            # Dispatch (sem Ex) reaproveita instÃ¢ncia jÃ¡ aberta e acelera conversÃµes seguintes.
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(str(src_docx), ReadOnly=True)
            # 17 = wdFormatPDF
            document.SaveAs(str(dst_pdf), FileFormat=17)
            document.Close(False)
            document = None
            if not keep_active:
                word.Quit()
                word = None

        if dst_pdf.exists():
            return True, None
        return False, "Word COM executou, mas nao gerou o PDF."
    except Exception as exc:
        return False, f"Word COM: {exc}"
    finally:
        try:
            if document is not None:
                document.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                if not keep_active:
                    word.Quit()
        except Exception:
            pass
        if com_ready:
            try:
                import pythoncom  # type: ignore
                pythoncom.CoUninitialize()
            except Exception:
                pass


def _convert_docx_to_pdf(src_docx: Path, dst_pdf: Path):
    # Serializa todo o pipeline de conversao para evitar corridas entre
    # Word COM, docx2pdf e LibreOffice em chamadas paralelas.
    with _WORD_CONVERT_LOCK:
        errors: list[str] = []
        logger.debug("Contrato: iniciando conversão DOCX->PDF (%s -> %s).", str(src_docx), str(dst_pdf))

        ok, err = _convert_via_word_com(src_docx, dst_pdf)
        if ok:
            logger.debug("Contrato: conversão concluída via Word COM.")
            return
        if err:
            errors.append(err)
            logger.warning("Contrato: Word COM falhou (%s).", err)

        try:
            from docx2pdf import convert as docx2pdf_convert  # type: ignore
            docx2pdf_convert(str(src_docx), str(dst_pdf))
            if dst_pdf.exists():
                logger.debug("Contrato: conversão concluída via docx2pdf.")
                return
            errors.append("docx2pdf executou, mas nao gerou o PDF.")
        except Exception as exc:
            errors.append(f"docx2pdf: {exc}")
            logger.warning("Contrato: docx2pdf falhou (%s).", exc)

        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            for candidate in (
                Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
                Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
            ):
                if candidate.exists():
                    soffice = str(candidate)
                    break

        if soffice:
            cmd = [
                str(soffice),
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(src_docx.parent),
                str(src_docx),
            ]
            proc = subprocess.run(cmd, capture_output=True, text=True)
            generated = src_docx.with_suffix(".pdf")
            if proc.returncode == 0 and generated.exists():
                try:
                    if generated.resolve() != dst_pdf.resolve():
                        if dst_pdf.exists():
                            dst_pdf.unlink()
                        generated.replace(dst_pdf)
                except Exception:
                    generated.replace(dst_pdf)
                logger.debug("Contrato: conversão concluída via LibreOffice.")
                return
            stderr = (proc.stderr or proc.stdout or "").strip()
            errors.append(f"LibreOffice: {stderr or 'falhou na conversao.'}")
            logger.warning("Contrato: LibreOffice falhou (%s).", stderr or "falhou na conversao.")
        else:
            errors.append("LibreOffice nao encontrado no sistema.")
            logger.warning("Contrato: LibreOffice não encontrado no sistema.")

        raise RuntimeError("Falha ao converter contrato para PDF. " + " | ".join(errors))
